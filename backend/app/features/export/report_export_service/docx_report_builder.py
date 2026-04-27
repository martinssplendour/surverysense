from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches as DocxInches
from docx.shared import Pt as DocxPt
from docx.shared import RGBColor

from app.features.export.report_export_service._constants import _REPORT_BODY_RGB, _REPORT_TITLE_RGB
from app.features.export.report_export_service.chart_image import DecodedChartImage


class DocxReportBuilder:
    def __init__(self, *, content_service) -> None:
        self.content_service = content_service

    def build(self, *, request, charts: list[DecodedChartImage]) -> bytes:
        document = Document()
        section = document.sections[0]
        section.top_margin = DocxInches(0.55)
        section.bottom_margin = DocxInches(0.55)
        section.left_margin = DocxInches(0.65)
        section.right_margin = DocxInches(0.65)

        title = document.add_paragraph()
        title.style = document.styles["Title"]
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.add_run(self.content_service.build_report_title())
        title_run.font.name = "Aptos Display"
        title_run.font.size = DocxPt(22)
        title_run.font.color.rgb = RGBColor(*_REPORT_TITLE_RGB)

        metadata = document.add_paragraph(self.content_service.build_subtitle(request))
        metadata.style = document.styles["Subtitle"]
        for run in metadata.runs:
            run.font.name = "Aptos"
            run.font.size = DocxPt(10)
            run.font.color.rgb = RGBColor(*_REPORT_TITLE_RGB)

        document.add_paragraph()

        if charts:
            for chart in charts:
                heading = document.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                heading_run = heading.add_run(chart.title)
                heading_run.bold = True
                heading_run.font.name = "Aptos Display"
                heading_run.font.size = DocxPt(14)
                heading_run.font.color.rgb = RGBColor(*_REPORT_TITLE_RGB)
                if chart.caption:
                    caption = document.add_paragraph(chart.caption)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        run.italic = True
                        run.font.name = "Aptos"
                        run.font.size = DocxPt(9)
                        run.font.color.rgb = RGBColor(*_REPORT_BODY_RGB)
                image_paragraph = document.add_paragraph()
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_paragraph.add_run().add_picture(BytesIO(chart.image_bytes), width=DocxInches(7.2))

        self._add_section_heading(document, self.content_service.build_summary_heading(request))
        group_sections = self.content_service.build_group_summary_sections(request)
        if group_sections:
            for section in group_sections[:8]:
                heading = document.add_paragraph()
                heading_run = heading.add_run(section.label)
                heading_run.bold = True
                heading_run.font.name = "Aptos"
                heading_run.font.size = DocxPt(11)
                heading_run.font.color.rgb = RGBColor(*_REPORT_TITLE_RGB)

                paragraph = document.add_paragraph()
                run = paragraph.add_run(section.summary)
                run.font.name = "Aptos"
                run.font.size = DocxPt(10)
                run.font.color.rgb = RGBColor(*_REPORT_BODY_RGB)
        else:
            for line in self.content_service.build_summary_lines(request):
                paragraph = document.add_paragraph()
                run = paragraph.add_run(line)
                run.font.name = "Aptos"
                run.font.size = DocxPt(10)
                run.font.color.rgb = RGBColor(*_REPORT_BODY_RGB)

        representative_sections = self.content_service.build_representative_sections(request)
        if representative_sections:
            self._add_section_heading(document, self.content_service.build_representative_heading())
            for label, examples in representative_sections:
                group_heading = document.add_paragraph()
                group_run = group_heading.add_run(label)
                group_run.bold = True
                group_run.font.name = "Aptos"
                group_run.font.size = DocxPt(11)
                group_run.font.color.rgb = RGBColor(*_REPORT_TITLE_RGB)
                for index, example in enumerate(examples, start=1):
                    paragraph = document.add_paragraph()
                    run = paragraph.add_run(f"{index}. {example}")
                    run.font.name = "Aptos"
                    run.font.size = DocxPt(10)
                    run.font.color.rgb = RGBColor(*_REPORT_BODY_RGB)

        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _add_section_heading(document: Document, text: str) -> None:
        heading = document.add_heading(level=1)
        heading_run = heading.add_run(text)
        heading_run.bold = True
        heading_run.font.name = "Aptos Display"
        heading_run.font.size = DocxPt(16)
        heading_run.font.color.rgb = RGBColor(*_REPORT_TITLE_RGB)
